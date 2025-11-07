# app.py — PIPE Demand Letter Service (Render-ready with CSV statement extraction)

from flask import Flask, request, send_file, abort, jsonify
from flask_cors import CORS
from io import BytesIO, StringIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
import pdfplumber
import csv
import os
import re

# -------------------------------------------------
# App & Config
# -------------------------------------------------
app = Flask(__name__)
CORS(app)  # allow Google Apps Script and other origins

API_KEY = os.getenv("PIPE_DEMAND_API_KEY", "YOUR_SECRET_KEY")
LETTERHEAD_IMAGE = os.getenv("PIPE_LETTERHEAD_IMAGE", "pipe_letterhead.png")  # optional file in repo root

# -------------------------------------------------
# Helpers
# -------------------------------------------------
CURRENCY_RE = re.compile(r"[^0-9.\-]")

def require_api_key(req) -> None:
    if req.headers.get("X-API-KEY") != API_KEY:
        abort(401, description="Unauthorized: Invalid API key.")

def norm_date(s: str, default: str = "") -> str:
    """Normalize many common date inputs to 'MMM DD, YYYY'."""
    if not s:
        return default
    s = str(s).strip()
    fmts = ("%b %d %Y", "%b %d, %Y", "%m %d %Y", "%m/%d/%Y", "%Y-%m-%d")
    for fmt in fmts:
        try:
            dt = datetime.strptime(s, fmt)
            return dt.strftime("%b %d, %Y")
        except Exception:
            pass
    return s

def parse_date_any(s):
    """Return datetime.date from many formats; else None."""
    if not s:
        return None
    s = str(s).strip()
    fmts = ("%b %d %Y", "%b %d, %Y", "%m %d %Y", "%m/%d/%Y", "%Y-%m-%d")
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt).date()
        except Exception:
            continue
    # try month abbreviations without zero-pad day (e.g., 'Nov 7 2025')
    try:
        return datetime.strptime(s.replace(",", ""), "%b %d %Y").date()
    except Exception:
        return None

def parse_money(val):
    """Accept '$12,345.67' or '12345.67' and return (float, '$12,345.67')."""
    if val in (None, ""):
        return None, ""
    try:
        f = float(CURRENCY_RE.sub("", str(val)))
        return f, f"${f:,.2f}"
    except Exception:
        return None, str(val)

def money(val) -> str:
    _, pretty = parse_money(val)
    return pretty

def normalize_rr(rr):
    """Accept '14', '14%', '14.0 foo' => '14%'."""
    if not rr:
        return ""
    m = re.search(r"(\d+(?:\.\d+)?)", str(rr))
    if not m:
        return str(rr)
    f = float(m.group(1))
    f = 0.0 if f < 0 else (100.0 if f > 100 else f)
    s = f"{f:.2f}".rstrip("0").rstrip(".")
    return s + "%"

def add_logo_if_present(doc: Document) -> None:
    """Insert letterhead logo top-left if file exists (non-fatal if missing)."""
    try:
        if os.path.isfile(LETTERHEAD_IMAGE):
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(LETTERHEAD_IMAGE, width=Inches(1.6))  # ~1.6" width
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass

def docx_to_text(stream: BytesIO):
    """Extract text from DOCX (for agreements uploaded as .docx)."""
    try:
        from docx import Document as DocxDocument
        d = DocxDocument(stream)
        return "\n".join((p.text or "") for p in d.paragraphs)
    except Exception:
        return None

def pdf_to_text(stream: BytesIO):
    """Extract text from text-based PDFs. Try pypdf, then pdfplumber."""
    text = []
    try:
        reader = PdfReader(stream)
        for p in reader.pages:
            text.append(p.extract_text() or "")
        joined = "\n".join(text)
        if joined.strip():
            return joined
    except Exception:
        pass
    try:
        stream.seek(0)
        with pdfplumber.open(stream) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)
    except Exception:
        return None

# -------------------------------------------------
# Routes: health & index
# -------------------------------------------------
@app.get("/")
def index():
    return {"status": "ok", "message": "Use POST /demand-letter (JSON), /agreement-extract (file), or /statement-extract (CSV)."}, 200

@app.get("/healthz")
def healthz():
    return {"status": "ok"}, 200

# -------------------------------------------------
# Route: Generate Demand Letter
# -------------------------------------------------
@app.post("/demand-letter")
def demand_letter():
    require_api_key(request)

    try:
        data = request.get_json(force=True) or {}
    except Exception:
        return jsonify({"error": "Invalid JSON body"}), 400

    # Inputs (accept both short and long names)
    business_name = data.get("business_name") or data.get("Business Name") or "BUSINESS NAME"
    business_address = data.get("business_address") or data.get("Business Address") or "Business address"
    contact_name = data.get("contact_name") or data.get("Contact Name") or "Client"

    today = norm_date(data.get("today") or data.get("Today") or datetime.utcnow().strftime("%b %d, %Y"))
    effective_date = norm_date(data.get("effective_date") or data.get("Effective Date"))
    default_date = norm_date(data.get("default_date") or data.get("Date of Default Event"))
    last_payment_date = norm_date(data.get("last_payment_date") or data.get("Date of Last Payment"))

    total_adv_plus_fee_raw = data.get("total_advance_plus_fee") or data.get("Total Advance + Fee")
    advance_amount_raw      = data.get("advance_amount")       or data.get("Advance Amount")
    fee_raw                 = data.get("fee")                  or data.get("Fee")
    total_revenue_raw       = data.get("total_revenue")        or data.get("Total Revenue Since Agreement to Today")
    rr_percent_raw          = data.get("rr_percent")           or data.get("Revenue Share Percentage (RR%)")
    rr_amount_raw           = data.get("rr_amount")            or data.get("Calculated % of Revenue Payable to Pipe ($)")
    successful_payments_raw = data.get("successful_payments")  or data.get("Amount of Successful Payments")
    percent_or_amount_due   = money(data.get("percent_or_amount_due") or data.get("Payment Percentage or Amount Due ($% of Revenue Amount)"))
    shortfall_raw           = data.get("shortfall")            or data.get("Shortfall")

    # Normalize
    total_adv_plus_fee = money(total_adv_plus_fee_raw)
    advance_amount     = money(advance_amount_raw)
    fee                = money(fee_raw)
    total_revenue      = money(total_revenue_raw)
    rr_percent         = normalize_rr(rr_percent_raw)

    _, rr_amount_money = parse_money(rr_amount_raw)
    _, successful_payments_money = parse_money(successful_payments_raw)
    _, shortfall_money = parse_money(shortfall_raw)

    # Auto-calc RR amount if blank
    if not rr_amount_money and total_revenue_raw and rr_percent:
        try:
            rr = float(re.search(r"(\d+(?:\.\d+)?)", rr_percent).group(1)) / 100.0
            tr = float(CURRENCY_RE.sub("", str(total_revenue_raw)))
            rr_calc = tr * rr
            rr_amount_money = f"${rr_calc:,.2f}"
        except Exception:
            pass

    # Auto-calc shortfall if blank
    if not shortfall_money and rr_amount_money and successful_payments_raw:
        try:
            rr_float = float(CURRENCY_RE.sub("", rr_amount_money))
            sp_float = float(CURRENCY_RE.sub("", successful_payments_raw))
            diff = max(0.0, rr_float - sp_float)
            shortfall_money = f"${diff:,.2f}"
        except Exception:
            pass

    # Build DOCX
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_logo_if_present(doc)  # top-left logo ~1.6"

    title = doc.add_paragraph("LETTER OF DEMAND")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph(f"{business_name}\n{business_address}\nUnited States of America\n")

    sent_p = doc.add_paragraph(f"SENT VIA EMAIL ON {today}")
    sent_p.runs[0].bold = True
    doc.add_paragraph("")

    re_p = doc.add_paragraph("Re: Demand for Payment - Pipe Merchant Cash Advance")
    re_p.runs[0].bold = True
    doc.add_paragraph("")

    dear = doc.add_paragraph(f"Dear {contact_name},")
    dear.runs[0].bold = True
    doc.add_paragraph("")

    body = (
        f"This is our last attempt and FINAL WARNING to seek payment for {business_name}'s merchant cash advance (\"MCA\") "
        f"before we seek all legal remedies available to us. {business_name} (\"you\") entered into an MCA Agreement "
        f"(\"Agreement\") with Pipe Advance LLC (the \"Company\") dated {effective_date} (the \"Effective Date\") for an MCA in "
        f"the total amount of {total_adv_plus_fee} (consisting of an MCA advance of {advance_amount} and a fee of {fee}).\n\n"

        f"Since {default_date}, {business_name} has failed to comply with its terms, by generating revenue and failing to "
        f"deliver and/or preventing Pipe from receiving its share of revenue payments. As of {today}, {business_name} has had "
        f"{total_revenue} in revenue payments of which {rr_percent} ({rr_amount_money or money(rr_amount_raw)}) are payable to Pipe "
        f"under the terms of the Agreement. We have only received {successful_payments_money or money(successful_payments_raw)} "
        f"towards your Total Advance Amount. The last payment to Pipe was on {last_payment_date}.\n\n"

        f"Your failure to pay Pipe the agreed upon percentage of revenue {percent_or_amount_due}, is a breach of the Agreement. "
        f"We have attempted to contact you and resolve this issue informally multiple times. Despite Pipe's continuous efforts to "
        f"resolve this issue, we have not received a payment.\n\n"

        f"If a payment of {shortfall_money or money(shortfall_raw)} is not received within 3 business days of receipt of this letter, "
        f"we will seek all remedies available to us under the Agreement, including referring this matter to a third-party collections "
        f"firm or seeking appropriate legal action. You may also be held liable and subject to additional fees incurred by Pipe in an "
        f"attempt to pursue these payments.\n\n"

        f"We urge you to treat this matter with the utmost urgency and to cooperate fully in resolving this breach amicably.\n\n"
    )
    doc.add_paragraph(body)

    doc.add_paragraph(
        "Please contact our Servicing and Collections Manager, William, at william@pipe.com immediately within the next 3 business days.\n\n"
        "Thank you for your immediate attention to this critical issue.\n\n"
        "Servicing and Collections\n"
        "Pipe Advance LLC"
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = re.sub(r"\s+", "_", business_name)
    fname = f"Demand_Letter_{safe_name}.docx"

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=fname
    )

# -------------------------------------------------
# Route: Agreement Extraction (kept from earlier)
# -------------------------------------------------
SUMMARY_BLOCK_RE = re.compile(
    r"(?:^|\n)\s*Pipe\s+Agreement\s*[\r\n]+Summary(.*?)(?:\nPayment\s*Method|\n\d+\s+pipe\.com|$)",
    re.IGNORECASE | re.DOTALL
)

def _first_group(pattern, text, flags=re.IGNORECASE | re.DOTALL):
    m = re.search(pattern, text, flags)
    return m.group(1).strip() if m else None

def _money_pretty(s):
    if not s:
        return None
    _, pretty = parse_money(s)
    return pretty or s

@app.post("/agreement-extract")
def agreement_extract():
    require_api_key(request)

    if "file" not in request.files:
        return jsonify({"ok": False, "error": "Missing file"}), 400
    f = request.files["file"]
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "Empty filename"}), 400

    name = f.filename.lower()
    stream = BytesIO(f.read()); stream.seek(0)

    # Read text (DOCX or PDF)
    text = None
    if name.endswith(".docx"):
        text = docx_to_text(stream)
    if not text:
        stream.seek(0)
        text = pdf_to_text(stream)

    if not text or not text.strip():
        return jsonify({"ok": False, "error": "Unable to read text (scanned PDF?)."}), 422

    # Prefer the "Summary" panel from page 1; fallback to entire doc
    m = SUMMARY_BLOCK_RE.search(text)
    panel = m.group(1) if m else text

    # Primary pulls from panel (tolerant to line breaks and spacing)
    out = {}

    # Business / Merchant
    out["business_name"] = (
        _first_group(r"(?:^|\n)\s*Merchant\s+([^\n]+)", panel) or
        _first_group(r"(?:^|\n)\s*Business\s*Name\s*[:\-–]\s*([^\n]+)", panel)
    )

    # Effective Date
    eff = (
        _first_group(r"(?:^|\n)\s*Effective\s*Date\s*[:\-–]?\s*([A-Za-z]{3,9}\s+\d{1,2}[,]?\s+\d{4}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})", panel) or
        _first_group(r"(?:^|\n)\s*Agreement\s*Date\s*[:\-–]?\s*([^\n]+)", panel)
    )
    out["effective_date"] = norm_date(eff) if eff else None

    # Amounts (allow descriptive text between label and $)
    adv = _first_group(r"Advance\s*Amount[\s\S]{0,160}?\$([\d,]+(?:\.\d{2})?)", panel)
    fee = _first_group(r"(?:^|\n)\s*Fee\b[\s\S]{0,160}?\$([\d,]+(?:\.\d{2})?)", panel)
    tot = _first_group(r"Total\s*Payment\s*Amount[\s\S]{0,200}?\$([\d,]+(?:\.\d{2})?)", panel)

    out["advance_amount"]         = _money_pretty(adv)
    out["fee"]                    = _money_pretty(fee)
    out["total_advance_plus_fee"] = _money_pretty(tot)

    # RR% (Payment/Remittance Rate)
    rr = _first_group(r"(?:Payment|Remittance|Withholding|Revenue\s*Share|RR%?)\s*Rate?[\s\S]{0,60}?(\d{1,2}(?:\.\d+)?%?)", panel)
    out["rr_percent"] = normalize_rr(rr) if rr else None

    # Partner / Processor (optional)
    out["partner"] = _first_group(r"(?:^|\n)\s*Partner\s*[^\n]*\n([A-Za-z0-9 &'.\-]+)", panel)

    # Fallbacks on whole document if still missing
    if not out.get("advance_amount"):
        adv2 = _first_group(r"Advance\s*Amount[\s\S]{0,200}?\$([\d,]+(?:\.\d{2})?)", text)
        out["advance_amount"] = _money_pretty(adv2) or out.get("advance_amount")
    if not out.get("fee"):
        fee2 = _first_group(r"(?:^|\n)\s*Fee\b[\s\S]{0,200}?\$([\d,]+(?:\.\d{2})?)", text)
        out["fee"] = _money_pretty(fee2) or out.get("fee")
    if not out.get("total_advance_plus_fee"):
        tot2 = _first_group(r"Total\s*(?:Payment|Advance|Purchase|Obligation)\s*Amount[\s\S]{0,240}?\$([\d,]+(?:\.\d{2})?)", text)
        out["total_advance_plus_fee"] = _money_pretty(tot2) or out.get("total_advance_plus_fee")
    if not out.get("rr_percent"):
        rr2 = _first_group(r"(?:Payment|Remittance|Withholding|Revenue\s*Share|RR%?)\s*Rate?[\s\S]{0,80}?(\d{1,2}(?:\.\d+)?%?)", text)
        out["rr_percent"] = normalize_rr(rr2) if rr2 else None
    if not out.get("business_name"):
        out["business_name"] = _first_group(r"(?:^|\n)\s*Merchant\s+([^\n]+)", text)

    # Final normalizations
    if out.get("effective_date"):
        out["effective_date"] = norm_date(out["effective_date"])

    print("EXTRACTED (agreement):", out)
    return jsonify({"ok": True, "extracted": out}), 200

# -------------------------------------------------
# NEW Route: Statement Extraction (CSV) — totals, due, paid, shortfall
# -------------------------------------------------
SUCCESS_STATUSES = {"succeeded", "paid", "completed", "posted", "captured", "success", "settled"}
FAIL_STATUSES    = {"returned", "failed", "declined", "chargeback", "void", "reversed"}

# Accept headers like the team uses (any case/extra spaces allowed):
# Revenue Date | Revenue | Collected | Method | Collection Date | Source | Increase | Status | External Link | Attempts
def _canonize(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", name.strip().lower())

COLUMN_ALIASES = {
    "revenue_date": {"revenue_date", "date", "payment_date"},
    "revenue": {"revenue", "gross_revenue", "amount"},
    "collected": {"collected", "pipe_collected", "to_pipe", "remitted"},
    "status": {"status", "result"},
    "collection_date": {"collection_date", "collected_date", "remit_date"},
    "method": {"method", "payment_method"},
}

def _find_col(mapping, key):
    want = COLUMN_ALIASES.get(key, {key})
    for k, v in mapping.items():
        if v in want:
            return k
    return None

@app.post("/statement-extract")
def statement_extract():
    """
    POST multipart/form-data:
      - file: CSV
      - effective_date: (string, optional) limit revenue from this date to today
      - rr_percent: (string, optional) to compute amount due
    Returns: totals and suggested fields to auto-fill in form.
    """
    require_api_key(request)

    if "file" not in request.files:
        return jsonify({"ok": False, "error": "Missing CSV file"}), 400
    f = request.files["file"]
    if not f or not f.filename.lower().endswith(".csv"):
        return jsonify({"ok": False, "error": "Upload a .csv statement"}), 400

    effective_date_str = request.form.get("effective_date", "").strip()
    rr_percent_in = request.form.get("rr_percent", "").strip()

    eff_date = parse_date_any(effective_date_str) if effective_date_str else None
    rr_norm = normalize_rr(rr_percent_in) if rr_percent_in else ""

    # Read CSV
    raw = f.read().decode("utf-8", errors="ignore")
    reader = csv.DictReader(StringIO(raw))
    if not reader.fieldnames:
        return jsonify({"ok": False, "error": "CSV missing header row"}), 400

    # Map headers
    mapping = {}
    for col in reader.fieldnames:
        mapping[col] = _canonize(col)

    # Reverse-map: canonical -> actual col name
    canon2actual = {v: k for k, v in mapping.items()}

    col_rev_date = _find_col(mapping, "revenue_date")
    col_revenue  = _find_col(mapping, "revenue")
    col_collected = _find_col(mapping, "collected")
    col_status   = _find_col(mapping, "status")

    # Default if not found
    if not col_revenue:
        return jsonify({"ok": False, "error": "CSV must include a 'Revenue' column"}), 400
    if not col_rev_date:
        return jsonify({"ok": False, "error": "CSV must include a 'Revenue Date' column"}), 400
    # collected/status are optional; we'll handle missing

    total_revenue = 0.0
    total_collected_success = 0.0

    for row in reader:
        try:
            # Date filter
            dt = parse_date_any(row.get(col_rev_date, "")) if col_rev_date else None
            if eff_date and dt and dt < eff_date:
                continue

            # Revenue
            rev_val = row.get(col_revenue, "")
            rev_float = float(CURRENCY_RE.sub("", str(rev_val))) if rev_val not in (None, "") else 0.0
            total_revenue += rev_float

            # Collected (only if success-like status)
            coll_val = row.get(col_collected, "") if col_collected else ""
            status_val = str(row.get(col_status, "") or "").strip().lower()
            status_canon = re.sub(r"[^a-z]", "", status_val)

            if coll_val not in (None, ""):
                coll_float = float(CURRENCY_RE.sub("", str(coll_val)))
            else:
                coll_float = 0.0

            if status_canon in SUCCESS_STATUSES or (col_status is None and coll_float > 0):
                total_collected_success += coll_float

        except Exception:
            # Skip any malformed lines
            continue

    # Compute RR amount due + shortfall if rr provided
    rr_amount_due = None
    shortfall = None
    if rr_norm:
        try:
            rr = float(re.search(r"(\d+(?:\.\d+)?)", rr_norm).group(1)) / 100.0
            rr_amount_due = total_revenue * rr
            shortfall = max(0.0, rr_amount_due - total_collected_success)
        except Exception:
            rr_norm = rr_norm  # leave as is

    resp = {
        "ok": True,
        "inputs_detected": {
            "revenue_rows_processed": reader.line_num - 1,
            "effective_date_applied": norm_date(effective_date_str) if eff_date else None,
            "columns": reader.fieldnames
        },
        "metrics": {
            "total_revenue": f"${total_revenue:,.2f}",
            "successful_payments": f"${total_collected_success:,.2f}",
            "rr_percent": rr_norm or None,
            "rr_amount": (f"${rr_amount_due:,.2f}" if rr_amount_due is not None else None),
            "shortfall": (f"${shortfall:,.2f}" if shortfall is not None else None),
        },
        # fields your form can auto-fill
        "fill": {
            "total_revenue": f"${total_revenue:,.2f}",
            "successful_payments": f"${total_collected_success:,.2f}",
            "rr_percent": rr_norm or "",
            "rr_amount": (f"${rr_amount_due:,.2f}" if rr_amount_due is not None else ""),
            "shortfall": (f"${shortfall:,.2f}" if shortfall is not None else "")
        }
    }
    print("EXTRACTED (statement):", resp["metrics"])
    return jsonify(resp), 200

# -------------------------------------------------
# Local run (optional)
# -------------------------------------------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080)
